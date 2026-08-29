import os
import re
import shutil
import glob
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

desktop_dirs = [
    r"C:\Users\Antonio\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Escritorio\Paymind Strategy"
]

for d in desktop_dirs:
    os.makedirs(d, exist_ok=True)

base_dir = r"C:\Users\Antonio\.gemini\antigravity-ide\scratch\paymind-growth-engine"
playbooks_dir = os.path.join(base_dir, "playbooks")
data_dir = os.path.join(base_dir, "data")

def style_heading(p, text, level):
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(37, 99, 235)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(30, 41, 59)

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except:
        with open(md_path, 'r', encoding='latin1') as f:
            lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        raw_line = line.rstrip('\n\r')
        stripped = raw_line.strip()

        if '|' in stripped and ('---' in stripped or stripped.startswith('|')):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table and '|' not in stripped:
            process_table(doc, table_lines)
            in_table = False
            table_lines = []

        if in_table:
            continue

        if not stripped:
            continue

        if stripped.startswith('# '):
            p = doc.add_paragraph()
            style_heading(p, stripped[2:].strip(), 1)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            style_heading(p, stripped[3:].strip(), 2)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            style_heading(p, stripped[4:].strip(), 3)
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Calibri'
            run.font.italic = True
            run.font.color.rgb = RGBColor(71, 85, 105)
        elif stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('1. ') or stripped.startswith('2. '):
            p = doc.add_paragraph(style='List Bullet' if stripped.startswith(('* ', '- ')) else 'List Number')
            p.paragraph_format.space_after = Pt(3)
            content = re.sub(r'^(\*|-|\d+\.)\s+', '', stripped)
            add_formatted_text(p, content)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, stripped)

    if in_table and table_lines:
        process_table(doc, table_lines)

    doc.save(docx_path)

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.font.name = 'Calibri'
            r.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(225, 29, 72)
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            link_text = part[1:part.index(']')]
            r = paragraph.add_run(link_text)
            r.font.name = 'Calibri'
            r.font.color.rgb = RGBColor(37, 99, 235)
            r.underline = True
        else:
            r = paragraph.add_run(part)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(51, 65, 85)

def process_table(doc, lines):
    rows_data = []
    for line in lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    num_rows = len(rows_data)
    num_cols = max(len(r) for r in rows_data)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                add_formatted_text(p, val)
                
                if r_idx == 0:
                    shd = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shd)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
                else:
                    if r_idx % 2 == 1:
                        shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shd)

    doc.add_paragraph()

# Buscar TODOS los archivos MD en playbooks y data
md_files = sorted(list(set(glob.glob(os.path.join(playbooks_dir, "*.md")) + glob.glob(os.path.join(data_dir, "*.md")))))

print(f"Encontrados {len(md_files)} archivos Markdown para convertir...")

converted_count = 0
for idx, src_md in enumerate(md_files, 1):
    fname = os.path.basename(src_md)
    clean_name = fname.replace('.md', '').replace('_', ' ').replace('-', ' ').title()
    clean_name = re.sub(r'[^\w\s]', '', clean_name).strip()
    docx_name = f"{idx:02d}_{clean_name}.docx"
    
    for d in desktop_dirs:
        if os.path.exists(d):
            target_docx = os.path.join(d, docx_name)
            parse_markdown_to_docx(src_md, target_docx)
            print(f"[OK] Creado Word ({idx}/{len(md_files)}): {os.path.basename(target_docx)}")
    converted_count += 1

# Copiar Excels, CSVs y HTMLs a la carpeta del Escritorio
all_data_files = glob.glob(os.path.join(data_dir, "*")) + glob.glob(os.path.join(playbooks_dir, "*.html"))

for src_file in all_data_files:
    if os.path.isfile(src_file) and not src_file.endswith('.md'):
        fname = os.path.basename(src_file)
        for d in desktop_dirs:
            if os.path.exists(d):
                dst_file = os.path.join(d, fname)
                shutil.copy2(src_file, dst_file)
                print(f"[OK] Copiado Asset: {fname} a {d}")

print(f"\nPROCESO COMPLETADO EXITOSAMENTE: {converted_count} documentos de Word (.docx) y todos los datos exportados a 'Paymind Strategy'.")
