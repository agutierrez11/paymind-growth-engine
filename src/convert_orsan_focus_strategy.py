import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

base_dir = r"C:\Users\Antonio\.gemini\antigravity-ide\scratch\paymind-growth-engine"
playbooks_dir = os.path.join(base_dir, "playbooks")
desktop_dirs = [
    r"C:\Users\Antonio\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Escritorio\Paymind Strategy"
]

src_md = os.path.join(playbooks_dir, "estrategia_foco_90_dias_caso_orsan_claude.md")
docx_name = "20_Estrategia_Foco_90_Dias_Caso_ORSAN_Claude.docx"

def style_heading(p, text, level):
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(37, 99, 235)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 41, 59)

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        raw = line.rstrip('\n\r').strip()
        if '|' in raw and ('---' in raw or raw.startswith('|')):
            in_table = True
            table_lines.append(raw)
            continue
        elif in_table and '|' not in raw:
            process_table(doc, table_lines)
            in_table = False
            table_lines = []

        if in_table or not raw:
            continue

        if raw.startswith('# '):
            p = doc.add_paragraph()
            style_heading(p, raw[2:], 1)
        elif raw.startswith('## '):
            p = doc.add_paragraph()
            style_heading(p, raw[3:], 2)
        elif raw.startswith('### '):
            p = doc.add_paragraph()
            style_heading(p, raw[4:], 3)
        elif raw.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(raw[2:])
            r.font.name = 'Calibri'
            r.font.italic = True
            r.font.color.rgb = RGBColor(71, 85, 105)
        elif raw.startswith('* ') or raw.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, raw[2:])
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, raw)

    if in_table and table_lines:
        process_table(doc, table_lines)

    doc.save(docx_path)

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.font.name = 'Calibri'
            r.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(225, 29, 72)
        else:
            r = paragraph.add_run(part)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

def process_table(doc, lines):
    rows_data = []
    for line in lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    table = doc.add_table(rows=len(rows_data), cols=max(len(r) for r in rows_data))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            if c_idx < len(row):
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                add_formatted_text(p, val)
                if r_idx == 0:
                    shd = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shd)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True

for d in desktop_dirs:
    if os.path.exists(d):
        dst = os.path.join(d, docx_name)
        parse_markdown_to_docx(src_md, dst)
        print(f"[OK] Creado documento Word de Estrategia de Foco Caso ORSAN: {dst}")

print("Conversión de Estrategia de Foco Caso ORSAN completada.")
