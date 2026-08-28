import os
import re
import shutil
import glob
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

base_dir = r"C:\Users\Antonio\.gemini\antigravity-ide\scratch\paymind-growth-engine"
playbooks_dir = os.path.join(base_dir, "playbooks")
desktop_dirs = [
    r"C:\Users\Antonio\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Escritorio\Paymind Strategy"
]

# 1. Reescribir el One-Pager Ejecutivo (Markdown y HTML) sin referencias a la agencia ni a la laptop
one_pager_md = """# 📄 RESUMEN EJECUTIVO: PLAN GTM Y FOCO A 90 DÍAS PAYMIND

> **Para:** Dirección General de PayMind  
> **De:** Antonio Gutiérrez | Consultor de Crecimiento & Pagos  
> **Asunto:** Plan de Aceleración GTM Gasolineras (Integración Volumétrica y Medición del Caso ORSAN)  
> **Tiempo de Lectura:** 2 Minutos

---

## ⚡ 1. Arquitectura del Plan GTM y Foco Operativo

El objetivo estratégico de PayMind en el sector gasolinero no es competir como un revendedor tradicional de terminales, sino consolidarse como la **Capa de Interoperabilidad e Infraestructura de Cobro** conectada directamente al sistema de control volumétrico (SAT Anexo 30) y a la adquirencia multibanco (BBVA, BanBajío, Afirme) con depósitos al día siguiente (T+1).

```
┌─────────────────────────────────────────────────────────────────────────────┐
│                    STACK DE INFRAESTRUCTURA DE PAGOS                        │
├─────────────────────────────────────────────────────────────────────────────┤
│ 1. VOLUMÉTRICO ➔ ControlGAS (ATIO), eGas, NexusFuel, Gasomarshal Odoo        │
│ 2. PAYMIND     ➔ Orquestación de Pagos Multi-Adquirente (T+1)                │
│ 3. HARDWARE    ➔ SmartPOS Nexgo ATEX (Inalámbrico antichispas en isla)      │
│ 4. OPERACIÓN   ➔ Conciliación Automática (Despacho - Pago - Factura)       │
└─────────────────────────────────────────────────────────────────────────────┘
```

---

## 🎯 2. Ejecución Prioritaria a 90 Días

Para garantizar tracción comercial inmediata y maximizar la eficiencia de recursos:

1. **Mes 1 (Auditoría Interna de Producción):** Extraer las métricas operativas y transaccionales del caso vivo **Grupo ORSAN (Mobil)** para disponer de evidencia empírica real de autorización, velocidad de despacho y tasa de adopción.
2. **Mes 2 (Partnership Técnico de Distribución):** Cerrar acuerdo de integración certificada con un proveedor de software volumétrico enfocado en independientes (**eGas o NexusFuel**).
3. **Mes 3 (Densidad en Clúster Regional):** Concentrar el esfuerzo comercial en el clúster geográfico de mayor densidad y acceso directo (**Sureste / Península**).

---

## 📊 3. Los 4 Activos de Prospección Listos para Despliegue

* **Base de Datos Segmentada (433 Gasolineras):** Clasificadas en Top Tier ABM, Mid Market e Independientes.
* **Secuencias de Prospección Socrática (Snov.io):** Mensajes directos de baja fricción enfocados en cobro al pie de bomba sin cambiar de banco adquirente.
* **Directorio ONEXPO (43 Asociaciones Estatales):** Mapeo de liderazgo gremial en los 32 estados.
* **Calculadora de Adquirencia Directa:** Herramienta interactiva para simular ahorros en tasa de descuento bancaria y flujo T+1.

---

## 💡 CONCLUSIÓN Y PRÓXIMOS PASOS

> **"El camino de menor fricción y mayor tracción para PayMind consiste en validar la línea base operativa con los datos reales de Grupo ORSAN, amarrar la integración técnica con un volumétrico de independientes y dominar la densidad regional en 90 días con evidencia empírica indiscutible."**
"""

one_pager_html = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Resumen Ejecutivo: Plan GTM PayMind</title>
    <style>
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #1e293b; max-width: 800px; margin: 40px auto; padding: 20px; background-color: #f8fafc; }
        .card { background: white; padding: 40px; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); border: 1px solid #e2e8f0; }
        h1 { color: #0f172a; border-bottom: 3px solid #2563eb; padding-bottom: 10px; font-size: 24px; }
        h2 { color: #2563eb; font-size: 18px; margin-top: 25px; }
        .meta { background: #eff6ff; padding: 15px; border-left: 4px solid #2563eb; border-radius: 4px; font-size: 14px; margin-bottom: 25px; }
        ul { padding-left: 20px; }
        li { margin-bottom: 8px; }
        .conclusion { background: #0f172a; color: white; padding: 20px; border-radius: 8px; margin-top: 30px; font-weight: 500; }
    </style>
</head>
<body>
    <div class="card">
        <h1>📄 RESUMEN EJECUTIVO: PLAN GTM PAYMIND</h1>
        <div class="meta">
            <strong>Para:</strong> Dirección General de PayMind<br>
            <strong>De:</strong> Antonio Gutiérrez | Consultor de Crecimiento & Pagos<br>
            <strong>Asunto:</strong> Plan de Aceleración GTM Gasolineras (Foco a 90 Días y Caso ORSAN)
        </div>

        <h2>⚡ 1. Arquitectura del Plan GTM</h2>
        <p>PayMind se posiciona como la <strong>Capa de Interoperabilidad e Infraestructura de Cobro</strong> conectada al control volumétrico (SAT Anexo 30) y a la adquirencia multibanco (BBVA, BanBajío, Afirme) con depósitos T+1.</p>

        <h2>🎯 2. Ejecución Prioritaria a 90 Días</h2>
        <ul>
            <li><strong>Mes 1:</strong> Auditoría interna de datos transaccionales del caso Grupo ORSAN (Mobil).</li>
            <li><strong>Mes 2:</strong> Alianza técnica con software volumétrico de independientes (eGas / NexusFuel).</li>
            <li><strong>Mes 3:</strong> Densidad comercial en clúster regional Sureste / Península.</li>
        </ul>

        <h2>📊 3. Activos Listos para Despliegue</h2>
        <ul>
            <li>Base de Datos Segmentada de 433 Gasolineras.</li>
            <li>Secuencias de correos socráticos para Snov.io.</li>
            <li>Directorio de 43 Asociaciones ONEXPO.</li>
            <li>Calculadora Web de Adquirencia Directa.</li>
        </ul>

        <div class="conclusion">
            "El camino de mayor tracción consiste en validar la línea base con el caso ORSAN, amarrar la integración técnica con un volumétrico independiente y dominar la densidad regional en 90 días."
        </div>
    </div>
</body>
</html>
"""

# Guardar MD y HTML limpios
with open(os.path.join(playbooks_dir, "ONE_PAGER_PLAN_GTM_VS_AGENCIA.md"), "w", encoding="utf-8") as f:
    f.write(one_pager_md)

with open(os.path.join(playbooks_dir, "ONE_PAGER_PLAN_GTM_VS_AGENCIA.html"), "w", encoding="utf-8") as f:
    f.write(one_pager_html)

# 2. Limpiar otros playbooks que contenían menciones a agencia/laptop
files_to_clean = [
    "propuesta_ejecutiva_ceo_paymind.md",
    "roadmap_comercial_y_caso_negocio.md",
    "plan_ajustes_anti_humo_zero_assumption_paymind.md"
]

for fname in files_to_clean:
    fpath = os.path.join(playbooks_dir, fname)
    if os.path.exists(fpath):
        with open(fpath, "r", encoding="utf-8") as f:
            content = f.read()
        # Reemplazar menciones a agencia o laptop
        content = re.sub(r'.*agencia.*', '', content, flags=re.IGNORECASE)
        content = re.sub(r'.*laptop.*', '', content, flags=re.IGNORECASE)
        content = re.sub(r'\n\s*\n', '\n\n', content)
        with open(fpath, "w", encoding="utf-8") as f:
            f.write(content)

# 3. Regenerar todos los documentos de Word (.docx) en el Escritorio
def style_heading(p, text, level):
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(37, 99, 235)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 41, 59)

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        raw = line.rstrip('\n\r').strip()
        if '|' in raw and ('---' in raw or raw.startswith('|')):
            in_table = True
            table_lines.append(raw)
            continue
        elif in_table and '|' not in raw:
            process_table(doc, table_lines)
            in_table = False
            table_lines = []

        if in_table or not raw:
            continue

        if raw.startswith('# '):
            p = doc.add_paragraph()
            style_heading(p, raw[2:], 1)
        elif raw.startswith('## '):
            p = doc.add_paragraph()
            style_heading(p, raw[3:], 2)
        elif raw.startswith('### '):
            p = doc.add_paragraph()
            style_heading(p, raw[4:], 3)
        elif raw.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(raw[2:])
            r.font.name = 'Calibri'
            r.font.italic = True
            r.font.color.rgb = RGBColor(71, 85, 105)
        elif raw.startswith('* ') or raw.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, raw[2:])
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, raw)

    if in_table and table_lines:
        process_table(doc, table_lines)

    doc.save(docx_path)

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.font.name = 'Calibri'
            r.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(225, 29, 72)
        else:
            r = paragraph.add_run(part)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

def process_table(doc, lines):
    rows_data = []
    for line in lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    table = doc.add_table(rows=len(rows_data), cols=max(len(r) for r in rows_data))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            if c_idx < len(row):
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                add_formatted_text(p, val)
                if r_idx == 0:
                    shd = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shd)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True

md_files = glob.glob(os.path.join(playbooks_dir, "*.md"))
for md in md_files:
    fname = os.path.basename(md)
    clean_name = fname.replace('.md', '').replace('_', ' ').title()
    docx_name = f"{clean_name}.docx"
    for d in desktop_dirs:
        if os.path.exists(d):
            dst = os.path.join(d, docx_name)
            parse_markdown_to_docx(md, dst)
            print(f"[OK] Re-generado Word limpio: {dst}")

print("Limpieza completa de referencias a agencia y laptop realizada con éxito.")
