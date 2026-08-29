import os
import re
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

base_dir = r"C:\Users\Antonio\.gemini\antigravity-ide\scratch\paymind-growth-engine"
playbooks_dir = os.path.join(base_dir, "playbooks")
data_dir = os.path.join(base_dir, "data")
desktop_dirs = [
    r"C:\Users\Antonio\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Escritorio\Paymind Strategy"
]

for d in desktop_dirs:
    os.makedirs(d, exist_ok=True)

# 1. Definición de las 4 Variantes A/B/C/D (Sin egocentrismo, sin clichés, sin asunciones)
md_ab_content = """# 🧪 Matriz de Pruebas A/B de Correos de Salida en Frío

> **Documento de Trabajo:** Variantes A/B/C/D para pruebas de tasa de apertura y respuesta en Snov.io.  
> **Regla de Oro:** Cero "soy especialista", cero "en empresas como la tuya", cero egocentrismo. Foco 100% en la acción práctica que beneficia a la estación.

---

## 🅰️ Variante A: Enfoque "Conserva tu Banco Actual" (Agnosticismo Financiero)

* **Hipótesis:** El gasolinero teme cambiar de banco o perder sus tasas. Ofrecer conectar la bomba a su banco actual genera alta apertura.

```text
Asunto: Conectar cobro en bomba manteniendo su banco actual

Hola,

Es posible conectar terminales inalámbricas al pie del auto directamente al dispensario, manteniendo la misma cuenta bancaria que ya utilizan (BBVA, BanBajío, Afirme) y recibiendo depósitos al día siguiente (T+1).

Si están abiertos a revisar cómo agilizar el cobro en pista sin cambiar de banco, ¿valdrá la pena platicar 5 minutos esta semana?

Saludos,
Antonio Gutiérrez
antonio.gutierrez@paymind.mx
```

---

## 🅱️ Variante B: Enfoque "Velocidad de Depósito T+1" (Liquidez)

* **Hipótesis:** El flujo de caja para la compra de pipas es prioritario. La promesa de dinero al día siguiente impulsa respuestas.

```text
Asunto: Depósito al día siguiente en cobros con tarjeta en bomba

Hola,

Ayudamos a que los cobros con tarjeta realizados al pie del dispensario se depositen al día siguiente (T+1), eliminando la espera en liquidaciones bancarias.

Las terminales son portátiles, inalámbricas y se conectan al sistema de la bomba respetando su número de afiliación actual.

Si les interesa evaluar una alternativa para acelerar el depósito de sus ventas, ¿tendrán 5 minutos esta semana?

Saludos,
Antonio Gutiérrez
antonio.gutierrez@paymind.mx
```

---

## 🅲 Variante C: Enfoque "Terminal Antichispas al Pie del Auto" (Operativo)

* **Hipótesis:** La durabilidad y seguridad del equipo en isla le importa al responsable operativo.

```text
Asunto: Terminales inalámbricas antichispas para cobro en isla

Hola,

Implementamos terminales portátiles SmartPOS Android con certificación antichispas (ATEX) diseñadas para uso rudo en la isla de carga.

El equipo se enlaza con el dispensario para cobrar el monto exacto, transmitiendo los datos sin requerir cables ni cambiar su banco adquirente.

¿Hará sentido platicar 5 minutos esta semana para revisar el equipo y la integración?

Saludos,
Antonio Gutiérrez
antonio.gutierrez@paymind.mx
```

---

## 🅳 Variante D: Enfoque Socrático Directo (Pregunta de Baja Fricción)

* **Hipótesis:** Una pregunta simple de 2 líneas sin rodeos genera mayor tasa de respuesta.

```text
Asunto: Consulta sobre cobro inalámbrico en dispensario

Hola,

¿Están evaluando conectar terminales inalámbricas directamente a sus bombas para cobrar tarjetas sin cambiar de banco adquirente?

Desarrollamos una pasarela que permite cobrar al pie del auto con depósitos al día siguiente (T+1).

Si el tema está en agenda, ¿valdrá la pena una llamada de 5 minutos esta semana?

Saludos,
Antonio Gutiérrez
antonio.gutierrez@paymind.mx
```
"""

# Guardar MD de pruebas A/B
ab_md_path = os.path.join(playbooks_dir, "variaciones_ab_testing_correos_frios.md")
with open(ab_md_path, "w", encoding="utf-8") as f:
    f.write(md_ab_content)

# Convertir a Word (.docx)
def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        raw = line.rstrip('\n\r').strip()
        if not raw:
            continue
        if raw.startswith('# '):
            p = doc.add_paragraph()
            r = p.add_run(raw[2:])
            r.bold = True
            r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif raw.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(raw[3:])
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(37, 99, 235)
        elif raw.startswith('* ') or raw.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(raw[2:])
        else:
            p = doc.add_paragraph()
            p.add_run(raw)
    doc.save(docx_path)

# Exportar a carpetas de Escritorio
docx_ab_name = "16_Variaciones_Ab_Testing_Correos_Frios.docx"
for d in desktop_dirs:
    target_docx = os.path.join(d, docx_ab_name)
    parse_markdown_to_docx(ab_md_path, target_docx)
    print(f"[OK] Creado Word A/B Testing: {target_docx}")

# 2. Actualizar CSV para Snov.io con las 4 Variantes A/B/C/D
csv_in = os.path.join(data_dir, "Snovio_Gasolineras_Segmentada_Clusters.csv")
df = pd.read_csv(csv_in, encoding='utf-8-sig')

df['Asunto_Variante_A'] = "Conectar cobro en bomba manteniendo su banco actual"
df['Cuerpo_Variante_A'] = (
    "Hola,\n\n"
    "Es posible conectar terminales inalámbricas al pie del auto directamente al dispensario, manteniendo la misma cuenta bancaria que ya utilizan (BBVA, BanBajío, Afirme) y recibiendo depósitos al día siguiente (T+1).\n\n"
    "Si están abiertos a revisar cómo agilizar el cobro en pista sin cambiar de banco, ¿valdrá la pena platicar 5 minutos esta semana?\n\n"
    "Saludos,\nAntonio Gutiérrez\nantonio.gutierrez@paymind.mx"
)

df['Asunto_Variante_B'] = "Depósito al día siguiente en cobros con tarjeta en bomba"
df['Cuerpo_Variante_B'] = (
    "Hola,\n\n"
    "Ayudamos a que los cobros con tarjeta realizados al pie del dispensario se depositen al día siguiente (T+1), eliminando la espera en liquidaciones bancarias.\n\n"
    "Las terminales son portátiles, inalámbricas y se conectan al sistema de la bomba respetando su número de afiliación actual.\n\n"
    "Si les interesa evaluar una alternativa para acelerar el depósito de sus ventas, ¿tendrán 5 minutos esta semana?\n\n"
    "Saludos,\nAntonio Gutiérrez\nantonio.gutierrez@paymind.mx"
)

df['Asunto_Variante_C'] = "Terminales inalámbricas antichispas para cobro en isla"
df['Cuerpo_Variante_C'] = (
    "Hola,\n\n"
    "Implementamos terminales portátiles SmartPOS Android con certificación antichispas (ATEX) diseñadas para uso rudo en la isla de carga.\n\n"
    "El equipo se enlaza con el dispensario para cobrar el monto exacto, transmitiendo los datos sin requerir cables ni cambiar su banco adquirente.\n\n"
    "¿Hará sentido platicar 5 minutos esta semana para revisar el equipo y la integración?\n\n"
    "Saludos,\nAntonio Gutiérrez\nantonio.gutierrez@paymind.mx"
)

df['Asunto_Variante_D'] = "Consulta sobre cobro inalámbrico en dispensario"
df['Cuerpo_Variante_D'] = (
    "Hola,\n\n"
    "¿Están evaluando conectar terminales inalámbricas directamente a sus bombas para cobrar tarjetas sin cambiar de banco adquirente?\n\n"
    "Desarrollamos una pasarela que permite cobrar al pie del auto con depósitos al día siguiente (T+1).\n\n"
    "Si el tema está en agenda, ¿valdrá la pena una llamada de 5 minutos esta semana?\n\n"
    "Saludos,\nAntonio Gutiérrez\nantonio.gutierrez@paymind.mx"
)

# Guardar CSV actualizado en data y en Escritorio
df.to_csv(csv_in, index=False, encoding='utf-8-sig')

for d in desktop_dirs:
    if os.path.exists(d):
        dst_csv = os.path.join(d, "Snovio_Gasolineras_Segmentada_Clusters.csv")
        df.to_csv(dst_csv, index=False, encoding='utf-8-sig')
        print(f"[OK] Actualizado CSV con Variantes A/B/C/D en: {dst_csv}")

print("Generación de A/B Testing completada.")
