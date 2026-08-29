import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

desktop_dirs = [
    r"C:\Users\Antonio\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Desktop\Paymind Strategy",
    r"C:\Users\Antonio\OneDrive\Escritorio\Paymind Strategy"
]

for d in desktop_dirs:
    os.makedirs(d, exist_ok=True)

doc = Document()

# Configurar margenes de 0.8 pulgadas
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def add_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(15, 23, 42)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(37, 99, 235)

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(30, 41, 59)

def add_box(subject, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.8)
    
    cell = table.cell(0, 0)
    
    shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)
    
    p_subj = cell.paragraphs[0]
    p_subj.paragraph_format.space_before = Pt(6)
    p_subj.paragraph_format.space_after = Pt(4)
    r_lbl = p_subj.add_run("📌 ASUNTO: ")
    r_lbl.bold = True
    r_lbl.font.name = 'Calibri'
    r_lbl.font.color.rgb = RGBColor(225, 29, 72)
    
    r_txt = p_subj.add_run(subject)
    r_txt.bold = True
    r_txt.font.name = 'Calibri'
    r_txt.font.size = Pt(12)
    r_txt.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sep = cell.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(4)
    r_sep = p_sep.add_run("✉️ CUERPO DEL CORREO (Copiar todo el bloque):")
    r_sep.bold = True
    r_sep.font.name = 'Calibri'
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(71, 85, 105)
    
    lines = body.strip().split('\n')
    for line in lines:
        p_body = cell.add_paragraph()
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(3)
        r_b = p_body.add_run(line)
        r_b.font.name = 'Calibri'
        r_b.font.size = Pt(11)
        r_b.font.color.rgb = RGBColor(30, 41, 59)
        
    doc.add_paragraph()

add_title("⚡ PRUEBA A/B ESPECIAL: CIERRE FISCAL DE MES (28 AL 31 DE AGOSTO)")

p_intro = doc.add_paragraph()
r_in = p_intro.add_run(
    "Esta campaña se enfoca 100% en el dolor operativo de fin de mes: la conciliación entre el cobro con tarjeta en pista y los reportes "
    "del control volumétrico para el cierre mensual. En Snov.io creas el Email 1 con prueba A/B (Variante A vs Variante B) para enviar HOY viernes 28, "
    "y el Follow-up para el LUNES 31 de agosto (Día del Cierre)."
)
r_in.font.name = 'Calibri'
r_in.font.italic = True
r_in.font.color.rgb = RGBColor(71, 85, 105)

# --- VARIANTE A ---
add_h2("🅰️ VARIANTE A: Ángulo Conciliación de Volumen en Pista")

add_h3("📅 ENVIAR HOY VIERNES 28 DE AGOSTO — Primer Contacto")
add_box(
    "Cierre de mes: conciliación de volumen y cobro en pista",
    "Buen día,\n\n"
    "Le escribo a su estación aprovechando el cierre de mes.\n\n"
    "Al cerrar agosto, ¿cómo resuelven la conciliación entre los pagos recibidos con tarjeta en la posición de carga y los volúmenes registrados en su sistema volumétrico?\n\n"
    "En PayMind trabajamos en la infraestructura de cobro en pista e integración directa con sistemas volumétricos (aclarando que no manejamos la emisión de facturación, sino la conectividad del cobro y la conciliación del volumen con el dispensario). Orsan aparece entre las empresas que confían públicamente en nuestra plataforma.\n\n"
    "¿Quién suele revisar esa parte durante el cierre en su estación?\n\n"
    "Saludos,\n"
    "Antonio Gutiérrez\n"
    "PayMind\n"
    "antonio.gutierrez@paymind.mx"
)

add_h3("📅 ENVIAR LUNES 31 DE AGOSTO — Día del Cierre Fiscal")
add_box(
    "Cuadre de volumen al cierre de mes",
    "Buen día,\n\n"
    "Hoy que concluye el mes, retomo la pregunta anterior.\n\n"
    "Cuando hay diferencias al cuadrar los cortes de caja entre las terminales de pista y los reportes del volumétrico, ¿cuánto tiempo toma a su equipo de administración u operaciones conciliar esas ventas?\n\n"
    "PayMind se conecta vía APIs con sistemas volumétricos para que la autorización del cobro en la posición quede ligada al registro de carga. Si no es un tema de su área, agradecería que me orientara con la persona adecuada.\n\n"
    "Saludos,\n"
    "Antonio Gutiérrez\n"
    "PayMind\n"
    "antonio.gutierrez@paymind.mx"
)

# --- VARIANTE B ---
add_h2("🅱️ VARIANTE B: Ángulo Enrutamiento al Responsable del Cierre Fiscal")

add_h3("📅 ENVIAR HOY VIERNES 28 DE AGOSTO — Primer Contacto")
add_box(
    "¿Quién lleva el cierre de volumen y pagos este fin de mes?",
    "Buen día,\n\n"
    "¿Podría orientarme sobre quién revisa la conciliación entre el cobro con tarjeta en la posición y el sistema volumétrico durante el cierre de mes?\n\n"
    "En PayMind trabajamos con infraestructura de pagos e integración directa con volumétricos (no manejamos facturación, sino el cobro en pista y la conciliación de volumen). Orsan aparece entre las empresas que confían públicamente en PayMind.\n\n"
    "No sé si corresponde a administración, operaciones o sistemas. Agradecería mucho que me indicara con quién conviene hablar.\n\n"
    "Saludos,\n"
    "Antonio Gutiérrez\n"
    "PayMind\n"
    "antonio.gutierrez@paymind.mx"
)

add_h3("📅 ENVIAR LUNES 31 DE AGOSTO — Día del Cierre Fiscal")
add_box(
    "¿Quién revisa el cuadre de ventas e inventario al cierre?",
    "Buen día,\n\n"
    "Aprovechando que hoy cierra el mes, busco evitar escribirle a la persona equivocada.\n\n"
    "¿Quién conoce la relación entre las terminales de pista, los comprobantes y la conciliación del reporte volumétrico mensual?\n\n"
    "Si no es un tema de su área, le agradecería cualquier orientación para dirigir el mensaje a la persona adecuada.\n\n"
    "Saludos,\n"
    "Antonio Gutiérrez\n"
    "PayMind\n"
    "antonio.gutierrez@paymind.mx"
)

out_filename = "00_CORREOS_FIN_DE_MES_A_B_TEST.docx"
for d in desktop_dirs:
    if os.path.exists(d):
        target_path = os.path.join(d, out_filename)
        doc.save(target_path)
        print(f"[OK] Creado documento Word de Cierre Fiscal A/B en: {target_path}")

print("Generación de Word Cierre Fiscal A/B completada exitosamente.")
